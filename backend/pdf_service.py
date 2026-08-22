"""Servicio de manipulacion de PDF para el flujo de Autocorreo.

Flujo de mesa:
  - Llega una simulacion (PDF) desde mesa (aprobaciones@centralmutuos.cl).
  - Se deja SOLO la pagina 1 (se eliminan plazos y gastos operacionales de la pag 2+).
  - El PDF ajustado se archiva por cliente y se puede reenviar.
"""
import io
import re
from pypdf import PdfReader, PdfWriter


def leer_texto(pdf_bytes, max_pages=2):
    """Extrae texto de las primeras paginas para clasificar el documento."""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        out = []
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            try:
                out.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(out)
    except Exception:
        return ""


def contar_paginas(pdf_bytes):
    try:
        return len(PdfReader(io.BytesIO(pdf_bytes)).pages)
    except Exception:
        return 0


def clasificar_documento(pdf_bytes, filename=""):
    """Devuelve 'simulacion', 'carta_aprobacion' u 'otro'.

    REGLA INVIOLABLE: cualquier indicio de que es una carta de aprobacion
    (nombre o contenido) la clasifica como carta y NUNCA se modifica.
    """
    fn = (filename or "").lower()
    if re.search(r"carta|aprobaci[oó]n|aprobacion", fn):
        return "carta_aprobacion"
    texto = (leer_texto(pdf_bytes) + " " + fn).lower()
    if re.search(r"agrado de informar|ha sido aprobad|carta de aprobaci|mutuo hipotecario.*aprobad", texto):
        return "carta_aprobacion"
    if re.search(r"simulaci[oó]n|dividendo|gastos operacionales|tasa|plazo|financiamiento|pie", texto):
        return "simulacion"
    return "otro"


def dejar_primera_pagina(pdf_bytes):
    """Genera un PDF nuevo con SOLO la primera pagina.

    Devuelve: (nuevos_bytes, paginas_originales, paginas_removidas)
    """
    reader = PdfReader(io.BytesIO(pdf_bytes))
    total = len(reader.pages)
    if total <= 1:
        return pdf_bytes, total, 0
    writer = PdfWriter()
    writer.add_page(reader.pages[0])
    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.getvalue(), total, total - 1


def extraer_texto_aprobacion(pdf_bytes):
    """Devuelve el parrafo de aprobacion si existe en el documento."""
    texto = leer_texto(pdf_bytes, max_pages=3)
    m = re.search(r"(estimad[oa][^.]*agrado de informar.*?)(?:\n\n|atentamente|saludos|$)",
                  texto, re.I | re.S)
    if m:
        return m.group(1).strip()
    return ""


IMG_EXT = (".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp", ".heic")

# PROTOCOLO 01-06 (REGLA INAMOVIBLE): toda imagen convertida a PDF recibe
# OBLIGATORIAMENTE un prefijo numérico de jerarquía antes de cualquier otra acción.
PREFIJOS_PROTOCOLO_IMG = [
    ("02_Liquidaciones", r"liquidaci|sueldo|remuneraci|haberes"),
    ("02_Impuesto_Renta", r"impuesto|renta|formulario 22|f22"),
    ("03_Certificado_AFP", r"afp|cotizaci|previred|afiliaci"),
    ("03_Resumen_Impuestos", r"boleta|honorario"),
    ("04_CMF", r"\bcmf\b|\bsmf\b|informe[_ ]?de[_ ]?deudas?|deuda"),
]


def prefijo_protocolo_imagen(base):
    """Imagen convertida = 01_Cedula_ por defecto (carnet), salvo que el
    nombre delate otra categoría del protocolo. Si ya trae prefijo, se respeta."""
    if re.match(r"^\d{2}_", base or ""):
        return base
    low = (base or "").lower()
    for pref, pat in PREFIJOS_PROTOCOLO_IMG:
        if re.search(pat, low):
            return f"{pref}_{base}"
    return f"01_Cedula_{base}"


RX_CEDULA_CONTENIDO = re.compile(r"c[eé]dula|identidad|registro civil|\brun\b", re.I)

# ══ DIVISOR DE PDF MULTI-DOCUMENTO (caso real: cliente escanea todo en un solo PDF) ══
CATS_PAGINA = [
    ("liquidacion", r"liquidaci[oó]n de sueldo|total haberes|l[ií]quido a pagar|sueldo base|remuneraci"),
    ("impuestos", r"impuesto a la renta|formulario 22|\bf22\b|declaraci[oó]n anual"),
    ("boletas", r"boleta.{0,25}honorario|honorarios electr[oó]nic"),
    ("afp", r"certificado de cotizaciones|cotizaciones previsionales|\bafp\b|previred"),
    ("cmf", r"comisi[oó]n para el mercado financiero|informe de deudas|\bcmf\b"),
    ("cedula", r"c[eé]dula|identidad|registro civ|\brun\b|apellidos"),
    ("contrato", r"contrato de trabajo|anexo de contrato"),
    ("subsidio", r"registro social de hogares|subsidio habitacional|serviu"),
]
PREF_CAT = {"cedula": "01_Cedula", "liquidacion": "02_Liquidaciones",
            "impuestos": "02_Impuesto_Renta", "boletas": "03_Resumen_Impuestos",
            "afp": "03_Certificado_AFP", "cmf": "04_CMF",
            "contrato": "05_Contrato", "subsidio": "06_Subsidio"}


def _categoria_pagina(texto):
    low = (texto or "").lower()
    for cat, pat in CATS_PAGINA:
        if re.search(pat, low):
            return cat
    return "otro"


def dividir_pdf_multidocumento(pdf_bytes, filename, max_paginas=30):
    """Detecta un PDF que contiene VARIOS documentos distintos y lo divide en PDFs
    separados por categoría de contenido (texto embebido o OCR por página).
    Devuelve lista [{filename, bytes, categoria, paginas}] o None si no aplica."""
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    n = len(reader.pages)
    if n < 2 or n > max_paginas:
        return None
    textos = [(p.extract_text() or "").strip() for p in reader.pages]
    faltantes = [i for i, t in enumerate(textos) if len(t) < 40]
    if faltantes:  # híbrido: OCR solo en las páginas escaneadas sin texto embebido
        try:
            from pdf2image import convert_from_bytes
            import ocr_service
            for i in faltantes:
                im = convert_from_bytes(pdf_bytes, dpi=200, first_page=i + 1, last_page=i + 1)[0]
                textos[i] = ocr_service.ocr_imagen(im)[0]
        except Exception:
            pass
    cats = [_categoria_pagina(t) for t in textos]
    for i in range(1, len(cats)):  # páginas sin señal heredan la del documento anterior
        if cats[i] == "otro":
            cats[i] = cats[i - 1]
    if len({c for c in cats if c != "otro"}) < 2:
        return None
    segmentos = []
    for i, c in enumerate(cats):
        if segmentos and segmentos[-1]["cat"] == c:
            segmentos[-1]["paginas"].append(i)
        else:
            segmentos.append({"cat": c, "paginas": [i]})
    base = re.sub(r"\.pdf$", "", filename or "documento", flags=re.I)
    base = re.sub(r"^\d{2}_[A-Za-z_]+?_", "", base)
    partes = []
    for idx, s in enumerate(segmentos, 1):
        w = pypdf.PdfWriter()
        for pi in s["paginas"]:
            w.add_page(reader.pages[pi])
        buf = io.BytesIO()
        w.write(buf)
        pref = PREF_CAT.get(s["cat"], "00_Documento")
        partes.append({"filename": f"{pref}_{base}_p{idx}.pdf", "bytes": buf.getvalue(),
                       "categoria": s["cat"], "paginas": [p + 1 for p in s["paginas"]]})
    return partes


def expandir_adjunto(raw_pdf, nombre):
    """[(bytes, nombre)] — si el PDF trae varios documentos, los separa; si no, lo deja igual."""
    try:
        partes = dividir_pdf_multidocumento(raw_pdf, nombre)
    except Exception:
        partes = None
    if partes:
        return [(p["bytes"], p["filename"]) for p in partes]
    return [(raw_pdf, nombre)]


def prefijo_por_contenido(base, texto_ocr=""):
    """PROTOCOLO 01-06 corregido: el prefijo se decide por el CONTENIDO real (OCR).
    El nombre original solo se usa como respaldo si la imagen no tiene texto legible."""
    if re.match(r"^\d{2}_", base or ""):
        return base
    cont = (texto_ocr or "").lower()
    if cont.strip():
        for pref, pat in PREFIJOS_PROTOCOLO_IMG:
            if re.search(pat, cont):
                return f"{pref}_{base}"
        if RX_CEDULA_CONTENIDO.search(cont):
            return f"01_Cedula_{base}"
    return prefijo_protocolo_imagen(base)


def convertir_a_pdf(raw_bytes, filename):
    """Convierte un archivo a PDF si no lo es.

    - PDF: se devuelve tal cual.
    - Imagenes (jpg/png/etc): se convierten con img2pdf/PIL.
    - Texto plano: se genera un PDF simple.
    Devuelve: (pdf_bytes, nuevo_nombre, convertido:bool). Lanza ValueError si no es soportado.
    """
    fn = (filename or "archivo").strip()
    low = fn.lower()
    if low.endswith(".pdf") or raw_bytes[:5] == b"%PDF-":
        return raw_bytes, fn, False
    base = fn.rsplit(".", 1)[0]
    if low.endswith(IMG_EXT):
        # FLUJO CORRECTO (orden del administrador): 1) preprocesar la imagen
        # (rotación OSD + escalado ≈300 DPI + contraste/nitidez), 2) OCR,
        # 3) clasificar por CONTENIDO real, 4) recién ahí renombrar con el prefijo.
        from PIL import Image
        import ocr_service as _ocr
        im, texto = None, ""
        try:
            im = Image.open(io.BytesIO(raw_bytes))
        except Exception:
            im = None
        if im is not None:
            texto, im = _ocr.ocr_imagen(im)
        nuevo = prefijo_por_contenido(base, texto) + ".pdf"
        if im is not None:
            jbuf = io.BytesIO()
            im.save(jbuf, format="JPEG", quality=92)
            try:
                import img2pdf
                return img2pdf.convert(jbuf.getvalue()), nuevo, True
            except Exception:
                out = io.BytesIO()
                im.save(out, format="PDF", resolution=300.0)
                return out.getvalue(), nuevo, True
        try:
            import img2pdf
            pdf = img2pdf.convert(raw_bytes)
            return pdf, nuevo, True
        except Exception:
            # Fallback PIL (convierte modos no soportados, ej. RGBA)
            im = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="PDF")
            return buf.getvalue(), nuevo, True
    if low.endswith((".txt", ".csv")):
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        y = 800
        for line in raw_bytes.decode("utf-8", errors="ignore").splitlines()[:120]:
            c.drawString(40, y, line[:110])
            y -= 14
            if y < 40:
                c.showPage(); y = 800
        c.showPage(); c.save(); buf.seek(0)
        return buf.getvalue(), base + ".pdf", True
    if low.endswith((".docx", ".doc", ".xlsx", ".xls")):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            lineas = []
            if low.endswith((".docx", ".doc")):
                import docx
                d = docx.Document(io.BytesIO(raw_bytes))
                for p in d.paragraphs:
                    if p.text.strip():
                        lineas.append(p.text)
                for t in d.tables:
                    for row in t.rows:
                        lineas.append(" | ".join(c.text.strip() for c in row.cells))
            else:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), data_only=True)
                for ws in wb.worksheets:
                    lineas.append(f"— Hoja: {ws.title} —")
                    for row in ws.iter_rows(values_only=True):
                        vals = [str(v) for v in row if v is not None]
                        if vals:
                            lineas.append(" | ".join(vals))
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=A4)
            y = 800
            for line in lineas[:600]:
                c.drawString(40, y, line[:110])
                y -= 14
                if y < 40:
                    c.showPage()
                    y = 800
            c.showPage()
            c.save()
            buf.seek(0)
            return buf.getvalue(), base + ".pdf", True
        except Exception as e:
            raise ValueError(f"No se pudo convertir el documento Office: {e}")
    raise ValueError(f"Formato no soportado para conversion: {fn}")


ROLES_TITULAR = re.compile(
    r"^(del?\s+|de\s+la\s+)?(cliente|asegurad|declarante|deudor|titular|"
    r"solicitante|representante|poderdante|mandante|apoderad)")
ROLES_CODEUDOR = re.compile(r"^(del?\s+|de\s+la\s+)?codeudor")


def posiciones_firma_cliente(pdf_bytes, rol="titular"):
    """Detecta las etiquetas donde debe firmar el cliente y devuelve sus posiciones
    [{pagina, x, top, alto_pagina, ancho_pagina, etiqueta}].

    Reglas (rol=titular, el cliente principal):
      - "Firma cliente/asegurado/declarante/deudor/titular/..." (NUNCA codeudor)
      - "Nombre y firma ..." (ej: declaración origen de fondos)
      - "Firma" sola en su línea (ej: designación de apoderado, PEP)
      - Se excluye "Firma ejecutivo ..."
    Con rol=codeudor solo se toman las etiquetas de codeudor."""
    import io as _io
    import unicodedata
    import pdfplumber

    def _norm(s):
        s = unicodedata.normalize("NFD", (s or "").lower())
        return "".join(ch for ch in s if unicodedata.category(ch) != "Mn")

    out = []
    with pdfplumber.open(_io.BytesIO(pdf_bytes)) as pdf:
        for pnum, page in enumerate(pdf.pages, start=1):
            try:
                words = page.extract_words() or []
            except Exception:
                try:
                    page.flush_cache()
                except Exception:
                    pass
                continue
            lineas = {}
            for w in words:
                lineas.setdefault(round(w["top"] / 4), []).append(w)
            for ws in lineas.values():
                ws.sort(key=lambda w: w["x0"])
                texto_linea = _norm(" ".join(x["text"] for x in ws))
                for i, w in enumerate(ws):
                    wt = _norm(w["text"])
                    if not wt.startswith("firma"):
                        continue
                    resto = _norm(" ".join(x["text"] for x in ws[i + 1:i + 4]))
                    prev = _norm(" ".join(x["text"] for x in ws[max(0, i - 2):i]))
                    es_codeudor_lbl = bool(ROLES_CODEUDOR.match(resto))
                    if rol == "codeudor":
                        ok = es_codeudor_lbl
                        etiqueta = "firma codeudor"
                    else:
                        if es_codeudor_lbl or resto.startswith("ejecutiv"):
                            continue
                        sola = re.fullmatch(r"firma[\s:_.]*", texto_linea) is not None
                        nombre_y = prev.endswith("nombre y")
                        con_rol = bool(ROLES_TITULAR.match(resto))
                        ok = con_rol or nombre_y or sola
                        etiqueta = ("firma " + resto.split(" ")[0]) if con_rol else \
                            ("nombre y firma" if nombre_y else "firma")
                    if ok:
                        out.append({"pagina": pnum, "x": float(w["x0"]),
                                    "top": float(w["top"]),
                                    "alto_pagina": float(page.height),
                                    "ancho_pagina": float(page.width),
                                    "etiqueta": etiqueta})
            # CIRUGÍA DE MEMORIA: liberar la página de la RAM apenas se extraen
            # sus coordenadas (evita que PDFs >5MB acumulen el layout completo)
            try:
                page.flush_cache()
                page.get_textmap.cache_clear()
            except Exception:
                pass
            del words, lineas
    return out


def estampar_referencias_firma(pdf_bytes, posiciones, nombre_firmante):
    """Imprime en cada etiqueta (menos la primera) una marca de referencia a la
    Firma Electrónica Avanzada única que cubre todo el documento (Ley 19.799)."""
    if not posiciones:
        return pdf_bytes
    from reportlab.pdfgen import canvas
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(io.BytesIO(pdf_bytes))
    por_pagina = {}
    for p in posiciones:
        por_pagina.setdefault(int(p["pagina"]), []).append(p)
    writer = PdfWriter()
    for i, page in enumerate(reader.pages, start=1):
        marcas = por_pagina.get(i)
        if marcas:
            w = float(page.mediabox.width)
            h = float(page.mediabox.height)
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=(w, h))
            for m in marcas:
                x = max(8, min(float(m["x"]), w - 210))
                y = h - float(m["top"]) + 3
                c.setFillColorRGB(0.13, 0.23, 0.42)
                c.setFont("Helvetica-Bold", 7.5)
                c.drawString(x, y + 17, "Firmado con Firma Electrónica Avanzada")
                c.setFont("Helvetica", 7.5)
                c.drawString(x, y + 8, (nombre_firmante or "").upper()[:48])
                c.setFont("Helvetica-Oblique", 6.3)
                c.drawString(x, y, "FEA e-CertChile válida para todo el documento (Ley 19.799)")
            c.save()
            buf.seek(0)
            overlay = PdfReader(buf).pages[0]
            page.merge_page(overlay)
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def estampar_pie_rastro(pdf_bytes, lineas):
    """Imprime un pie de rastro (trazabilidad de firma) en todas las páginas."""
    from reportlab.pdfgen import canvas
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=(w, h))
        c.setFillColorRGB(0.32, 0.32, 0.32)
        y = 7
        for ln in reversed(lineas):
            c.setFont("Helvetica", 5.6)
            c.drawString(14, y, ln[:170])
            y += 7
        c.save()
        buf.seek(0)
        page.merge_page(PdfReader(buf).pages[0])
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()
